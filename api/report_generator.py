from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import sys
import pdfplumber
from openai import OpenAI
from datetime import datetime
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Pinecone as PineconeVectorStore
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from io import BytesIO
from docx import Document
from docx.shared import Pt
from pinecone import Pinecone
import base64
import json

# 환경 변수 로딩
load_dotenv()

app = Flask(__name__)
CORS(app)

# OpenAI 클라이언트 설정
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

# Pinecone 설정
pinecone_api_key = os.getenv("PINECONE_API_KEY")
pinecone_env = os.getenv("PINECONE_ENV")
pc = Pinecone(api_key=pinecone_api_key)
index_name = "carbone-index"

# PDF 텍스트 추출
def extract_text_from_pdf(pdf_data):
    """PDF 데이터에서 텍스트를 추출합니다."""
    try:
        with pdfplumber.open(BytesIO(pdf_data)) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        return f"PDF 파싱 오류: {str(e)}"

# 목차 추출 함수
def extract_table_of_contents(text):
    """문서에서 목차를 추출합니다."""
    try:
        prompt = f"""
다음 문서에서 **목차(차례)**에 해당하는 부분만 정확히 추출해 주세요.
- 숫자나 로마자, 제목 패턴을 이용해 목차 항목만 뽑아주세요.
- 본문 내용은 포함하지 말고, 목차 구조만 출력하세요.

문서 내용:
{text[:4000]}
"""
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "당신은 문서 구조에서 목차만 정확히 추출하는 AI입니다."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"목차 추출 오류: {str(e)}"

# 문서 전체 구조 요약
def summarize_template_structure(text):
    """문서의 형식을 분석하고 요약합니다."""
    try:
        prompt = f"""
다음 문서의 형식(보고서 구조, 제목 스타일, 구성 흐름 등)을 간단히 요약해 주세요.
- 문서가 어떤 형식으로 작성되어 있는지 설명해주세요.
- 목차, 본문 구성, 언어 톤 등을 포함해 형식을 분석해주세요.

문서 내용:
{text[:4000]}
"""
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "당신은 문서 형식을 분석하고 요약하는 AI입니다."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"문서 구조 분석 오류: {str(e)}"

# Pinecone 기반 RAG용 DB 생성
def create_vector_store(text):
    """텍스트를 벡터 스토어에 저장합니다."""
    try:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        docs = text_splitter.create_documents([text])
        embeddings = OpenAIEmbeddings()

        # 인덱스 존재 여부 확인 → 없으면 생성
        if index_name not in pc.list_indexes().names():
            from pinecone import ServerlessSpec
            pc.create_index(
                name=index_name,
                dimension=1536,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )

        vector_store = PineconeVectorStore.from_documents(docs, embeddings, index_name=index_name)
        return vector_store
    except Exception as e:
        return f"벡터 스토어 생성 오류: {str(e)}"

# 보고서 생성
def generate_report_with_rag(topic, vector_store):
    """RAG를 사용하여 보고서를 생성합니다."""
    try:
        retriever = vector_store.as_retriever(search_kwargs={"k": 5})
        llm = ChatOpenAI(temperature=0.7, model_name="gpt-4")
        qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)

        prompt = f"""
'{topic}'에 대한 보고서를 작성해 주세요. 
- 참고 문서를 기반으로 구성과 형식을 따라주세요.
- 목차 구성, 말투, 분석 방향을 유지하고, 새로운 주제에 맞게 내용을 작성하세요.
- 참고 문서는 참고만 할뿐이지, 거기의 내용을 가지고와서는 안됩니다.
"""
        return qa.run(prompt)
    except Exception as e:
        return f"보고서 생성 오류: {str(e)}"

# Word 보고서 생성 함수
def generate_docx_report(text, topic):
    """Word 문서를 생성합니다."""
    try:
        doc = Document()
        doc.add_heading(topic, level=1)

        for paragraph in text.split("\n"):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                p.style.font.size = Pt(11)
                p.paragraph_format.line_spacing = 1.5

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')
    except Exception as e:
        return f"Word 문서 생성 오류: {str(e)}"

@app.route('/api/analyze-document', methods=['POST'])
def analyze_document():
    """PDF 문서를 분석합니다."""
    try:
        data = request.get_json()
        pdf_base64 = data.get('pdf_data')
        
        if not pdf_base64:
            return jsonify({'error': 'PDF 데이터가 없습니다.'}), 400

        # Base64 디코딩
        pdf_data = base64.b64decode(pdf_base64.split(',')[1] if ',' in pdf_base64 else pdf_base64)
        
        # PDF 텍스트 추출
        extracted_text = extract_text_from_pdf(pdf_data)
        
        if extracted_text.startswith("PDF 파싱 오류"):
            return jsonify({'error': extracted_text}), 400

        # 목차 추출
        toc = extract_table_of_contents(extracted_text)
        
        # 문서 구조 요약
        structure_summary = summarize_template_structure(extracted_text)
        
        # 벡터 스토어 생성
        vector_store = create_vector_store(extracted_text)
        
        return jsonify({
            'success': True,
            'toc': toc,
            'structure': structure_summary,
            'text_length': len(extracted_text)
        })
        
    except Exception as e:
        return jsonify({'error': f'서버 오류: {str(e)}'}), 500

@app.route('/api/generate-report', methods=['POST'])
def generate_report():
    """보고서를 생성합니다."""
    try:
        data = request.get_json()
        topic = data.get('topic')
        pdf_base64 = data.get('pdf_data')
        
        if not topic or not pdf_base64:
            return jsonify({'error': '주제와 PDF 데이터가 필요합니다.'}), 400

        # Base64 디코딩
        pdf_data = base64.b64decode(pdf_base64.split(',')[1] if ',' in pdf_base64 else pdf_base64)
        
        # PDF 텍스트 추출
        extracted_text = extract_text_from_pdf(pdf_data)
        
        if extracted_text.startswith("PDF 파싱 오류"):
            return jsonify({'error': extracted_text}), 400

        # 벡터 스토어 생성
        vector_store = create_vector_store(extracted_text)
        
        # 보고서 생성
        generated_report = generate_report_with_rag(topic, vector_store)
        
        if generated_report.startswith("보고서 생성 오류"):
            return jsonify({'error': generated_report}), 400

        # Word 문서 생성
        docx_base64 = generate_docx_report(generated_report, topic)
        
        return jsonify({
            'success': True,
            'report': generated_report,
            'docx_data': docx_base64
        })
        
    except Exception as e:
        return jsonify({'error': f'서버 오류: {str(e)}'}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """서버 상태 확인"""
    return jsonify({'status': 'healthy', 'timestamp': datetime.now().isoformat()})

if __name__ == '__main__':
    app.run(debug=True, port=5001) 